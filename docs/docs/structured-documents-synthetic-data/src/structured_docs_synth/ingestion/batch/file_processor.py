#!/usr/bin/env python3
"""
File Processor for batch file ingestion and processing
"""

import os
import time
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, Iterator
from dataclasses import dataclass, field
from enum import Enum
import hashlib

from pydantic import BaseModel, Field

from ...core.config import get_config
from ...core.logging import get_logger
from ...core.exceptions import ProcessingError, ValidationError


class FileType(Enum):
    """Supported file types"""
    PDF = "pdf"
    IMAGE = "image"
    DOC = "doc"
    DOCX = "docx"
    TXT = "txt"
    JSON = "json"
    XML = "xml"
    HTML = "html"
    UNKNOWN = "unknown"


class ProcessingStatus(Enum):
    """File processing status"""
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    SKIPPED = "skipped"


@dataclass
class FileMetadata:
    """File metadata information"""
    file_path: str
    file_name: str
    file_size: int
    file_type: FileType
    mime_type: str
    hash_md5: str
    created_time: float
    modified_time: float
    encoding: Optional[str] = None
    page_count: Optional[int] = None
    dimensions: Optional[tuple] = None


@dataclass
class ProcessedFile:
    """Processed file result"""
    metadata: FileMetadata
    status: ProcessingStatus
    extracted_text: str = ""
    extracted_data: Dict[str, Any] = field(default_factory=dict)
    error_message: str = ""
    processing_time: float = 0.0
    artifacts: List[str] = field(default_factory=list)  # Generated files


class FileProcessorConfig(BaseModel):
    """File processor configuration"""
    
    # Processing settings
    max_file_size_mb: float = Field(100.0, description="Maximum file size in MB")
    max_files_per_batch: int = Field(1000, description="Maximum files per batch")
    parallel_workers: int = Field(4, description="Number of parallel workers")
    
    # File filtering
    allowed_extensions: List[str] = Field(
        default=[".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".txt", ".docx", ".json"],
        description="Allowed file extensions"
    )
    excluded_patterns: List[str] = Field(
        default=[".*", "__*", "*.tmp"],
        description="Excluded file patterns"
    )
    
    # Text extraction
    extract_text: bool = Field(True, description="Extract text from documents")
    extract_metadata: bool = Field(True, description="Extract file metadata")
    preserve_formatting: bool = Field(False, description="Preserve text formatting")
    
    # Output settings
    output_dir: Optional[str] = Field(None, description="Output directory for artifacts")
    create_thumbnails: bool = Field(False, description="Create image thumbnails")
    thumbnail_size: int = Field(256, description="Thumbnail size in pixels")
    
    # Error handling
    skip_errors: bool = Field(True, description="Skip files with errors")
    max_retries: int = Field(3, description="Maximum retry attempts")


class FileProcessor:
    """
    File Processor for batch file ingestion
    
    Handles various file types, extracts content, metadata, and manages
    batch processing with error handling and progress tracking.
    """
    
    def __init__(self, config: Optional[FileProcessorConfig] = None):
        self.logger = get_logger(__name__)
        self.config = config or FileProcessorConfig()
        
        # Setup output directory
        if self.config.output_dir:
            self.output_dir = Path(self.config.output_dir)
            self.output_dir.mkdir(parents=True, exist_ok=True)
        else:
            self.output_dir = None
        
        self.logger.info("File Processor initialized")
    
    def process_files(self, source_path: Union[str, Path], 
                     recursive: bool = True) -> List[ProcessedFile]:
        """Process files from source directory or file list"""
        start_time = time.time()
        source_path = Path(source_path)
        
        try:
            # Discover files
            file_paths = self._discover_files(source_path, recursive)
            
            self.logger.info(f"Found {len(file_paths)} files to process")
            
            # Filter files
            filtered_paths = self._filter_files(file_paths)
            
            self.logger.info(f"Processing {len(filtered_paths)} files after filtering")
            
            # Process files
            results = []
            for i, file_path in enumerate(filtered_paths):
                try:
                    result = self._process_single_file(file_path)
                    results.append(result)
                    
                    if (i + 1) % 100 == 0:
                        self.logger.info(f"Processed {i + 1}/{len(filtered_paths)} files")
                        
                except Exception as e:
                    if not self.config.skip_errors:
                        raise
                    
                    error_result = ProcessedFile(
                        metadata=self._extract_basic_metadata(file_path),
                        status=ProcessingStatus.FAILED,
                        error_message=str(e)
                    )
                    results.append(error_result)
                    self.logger.warning(f"Failed to process {file_path}: {e}")
            
            processing_time = time.time() - start_time
            
            # Log summary
            successful = sum(1 for r in results if r.status == ProcessingStatus.COMPLETED)
            failed = sum(1 for r in results if r.status == ProcessingStatus.FAILED)
            
            self.logger.info(f"File processing completed in {processing_time:.2f}s. "
                           f"Success: {successful}, Failed: {failed}")
            
            return results
            
        except Exception as e:
            self.logger.error(f"File processing failed: {e}")
            raise ProcessingError(f"File processing error: {e}")
    
    def _discover_files(self, source_path: Path, recursive: bool) -> List[Path]:
        """Discover files from source path"""
        files = []
        
        if source_path.is_file():
            files.append(source_path)
        elif source_path.is_dir():
            if recursive:
                files.extend(source_path.rglob("*"))
            else:
                files.extend(source_path.iterdir())
            
            # Filter to files only
            files = [f for f in files if f.is_file()]
        else:
            raise ValueError(f"Source path does not exist: {source_path}")
        
        return files
    
    def _filter_files(self, file_paths: List[Path]) -> List[Path]:
        """Filter files based on configuration"""
        filtered = []
        
        for file_path in file_paths:
            # Check extension
            if file_path.suffix.lower() not in self.config.allowed_extensions:
                continue
            
            # Check excluded patterns
            if any(file_path.match(pattern) for pattern in self.config.excluded_patterns):
                continue
            
            # Check file size
            try:
                file_size = file_path.stat().st_size
                max_size = self.config.max_file_size_mb * 1024 * 1024
                if file_size > max_size:
                    self.logger.warning(f"Skipping large file {file_path} ({file_size} bytes)")
                    continue
            except OSError:
                continue
            
            filtered.append(file_path)
            
            # Check max files limit
            if len(filtered) >= self.config.max_files_per_batch:
                self.logger.warning(f"Reached maximum files limit ({self.config.max_files_per_batch})")
                break
        
        return filtered
    
    def _process_single_file(self, file_path: Path) -> ProcessedFile:
        """Process a single file"""
        start_time = time.time()
        
        # Extract metadata
        metadata = self._extract_file_metadata(file_path)
        
        # Initialize result
        result = ProcessedFile(
            metadata=metadata,
            status=ProcessingStatus.PROCESSING
        )
        
        try:
            # Extract text if enabled
            if self.config.extract_text:
                result.extracted_text = self._extract_text(file_path, metadata.file_type)
            
            # Extract structured data
            result.extracted_data = self._extract_structured_data(file_path, metadata.file_type)
            
            # Create artifacts if needed
            if self.output_dir:
                result.artifacts = self._create_artifacts(file_path, metadata)
            
            result.status = ProcessingStatus.COMPLETED
            
        except Exception as e:
            result.status = ProcessingStatus.FAILED
            result.error_message = str(e)
            self.logger.debug(f"File processing error for {file_path}: {e}")
        
        result.processing_time = time.time() - start_time
        return result
    
    def _extract_file_metadata(self, file_path: Path) -> FileMetadata:
        """Extract comprehensive file metadata"""
        stat = file_path.stat()
        
        # Calculate file hash
        hash_md5 = self._calculate_file_hash(file_path)
        
        # Detect file type
        file_type = self._detect_file_type(file_path)
        
        # Get MIME type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        mime_type = mime_type or "application/octet-stream"
        
        metadata = FileMetadata(
            file_path=str(file_path),
            file_name=file_path.name,
            file_size=stat.st_size,
            file_type=file_type,
            mime_type=mime_type,
            hash_md5=hash_md5,
            created_time=stat.st_ctime,
            modified_time=stat.st_mtime
        )
        
        # Extract additional metadata based on file type
        if file_type == FileType.IMAGE:
            metadata.dimensions = self._get_image_dimensions(file_path)
        elif file_type in [FileType.PDF, FileType.DOC, FileType.DOCX]:
            metadata.page_count = self._get_document_page_count(file_path)
        
        return metadata
    
    def _extract_basic_metadata(self, file_path: Path) -> FileMetadata:
        """Extract basic metadata for failed files"""
        try:
            stat = file_path.stat()
            return FileMetadata(
                file_path=str(file_path),
                file_name=file_path.name,
                file_size=stat.st_size,
                file_type=self._detect_file_type(file_path),
                mime_type="unknown",
                hash_md5="",
                created_time=stat.st_ctime,
                modified_time=stat.st_mtime
            )
        except Exception:
            return FileMetadata(
                file_path=str(file_path),
                file_name=file_path.name,
                file_size=0,
                file_type=FileType.UNKNOWN,
                mime_type="unknown",
                hash_md5="",
                created_time=0,
                modified_time=0
            )
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate MD5 hash of file"""
        try:
            hasher = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception:
            return ""
    
    def _detect_file_type(self, file_path: Path) -> FileType:
        """Detect file type from extension and content"""
        extension = file_path.suffix.lower()
        
        if extension == ".pdf":
            return FileType.PDF
        elif extension in [".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"]:
            return FileType.IMAGE
        elif extension == ".doc":
            return FileType.DOC
        elif extension == ".docx":
            return FileType.DOCX
        elif extension == ".txt":
            return FileType.TXT
        elif extension == ".json":
            return FileType.JSON
        elif extension in [".xml", ".xhtml"]:
            return FileType.XML
        elif extension in [".html", ".htm"]:
            return FileType.HTML
        else:
            return FileType.UNKNOWN
    
    def _extract_text(self, file_path: Path, file_type: FileType) -> str:
        """Extract text from file based on type"""
        try:
            if file_type == FileType.TXT:
                return self._extract_text_from_txt(file_path)
            elif file_type == FileType.PDF:
                return self._extract_text_from_pdf(file_path)
            elif file_type == FileType.DOCX:
                return self._extract_text_from_docx(file_path)
            elif file_type == FileType.JSON:
                return self._extract_text_from_json(file_path)
            elif file_type == FileType.XML:
                return self._extract_text_from_xml(file_path)
            elif file_type == FileType.HTML:
                return self._extract_text_from_html(file_path)
            else:
                return ""
        except Exception as e:
            self.logger.debug(f"Text extraction failed for {file_path}: {e}")
            return ""
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF"""
        try:
            # Try PyPDF2 first
            try:
                import PyPDF2
                text_parts = []
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        text_parts.append(page.extract_text())
                return '\n'.join(text_parts)
            except ImportError:
                pass
            
            # Try pdfplumber as fallback
            try:
                import pdfplumber
                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                return '\n'.join(text_parts)
            except ImportError:
                pass
            
            # Try PyMuPDF as last resort
            try:
                import fitz  # PyMuPDF
                text_parts = []
                doc = fitz.open(file_path)
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return '\n'.join(text_parts)
            except ImportError:
                pass
            
            # If no PDF library available, return placeholder
            self.logger.warning(f"No PDF library available for {file_path}. Install PyPDF2, pdfplumber, or PyMuPDF.")
            return f"[PDF content from {file_path.name} - no PDF library available]"
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            return f"[Failed to extract PDF content from {file_path.name}: {str(e)}]"
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX"""
        try:
            import docx
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
            
        except ImportError:
            self.logger.warning(f"python-docx not installed. Cannot extract text from {file_path}")
            return f"[DOCX content from {file_path.name} - python-docx not installed]"
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return f"[Failed to extract DOCX content from {file_path.name}: {str(e)}]"
    
    def _extract_text_from_json(self, file_path: Path) -> str:
        """Extract text from JSON file"""
        import json
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            # Extract text values from JSON
            return self._extract_text_from_json_recursive(data)
    
    def _extract_text_from_json_recursive(self, obj) -> str:
        """Recursively extract text from JSON object"""
        if isinstance(obj, str):
            return obj + " "
        elif isinstance(obj, dict):
            return "".join(self._extract_text_from_json_recursive(v) for v in obj.values())
        elif isinstance(obj, list):
            return "".join(self._extract_text_from_json_recursive(item) for item in obj)
        else:
            return str(obj) + " "
    
    def _extract_text_from_xml(self, file_path: Path) -> str:
        """Extract text from XML file"""
        try:
            # Try BeautifulSoup first
            try:
                from bs4 import BeautifulSoup
                with open(file_path, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f, 'xml')
                    # Remove script and style elements
                    for script in soup(['script', 'style']):
                        script.decompose()
                    return soup.get_text(separator=' ', strip=True)
            except ImportError:
                pass
            
            # Try lxml as fallback
            try:
                from lxml import etree
                tree = etree.parse(str(file_path))
                return ' '.join(tree.xpath('//text()'))
            except ImportError:
                pass
            
            # Use basic xml.etree as last resort
            import xml.etree.ElementTree as ET
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            def extract_text(element):
                text_parts = []
                if element.text:
                    text_parts.append(element.text.strip())
                for child in element:
                    text_parts.extend(extract_text(child))
                if element.tail:
                    text_parts.append(element.tail.strip())
                return text_parts
            
            return ' '.join(extract_text(root))
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from XML {file_path}: {e}")
            return f"[Failed to extract XML content from {file_path.name}: {str(e)}]"
    
    def _extract_text_from_html(self, file_path: Path) -> str:
        """Extract text from HTML file"""
        try:
            # Try BeautifulSoup
            try:
                from bs4 import BeautifulSoup
                with open(file_path, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f, 'html.parser')
                    
                    # Remove script and style elements
                    for script in soup(['script', 'style']):
                        script.decompose()
                    
                    # Get text and preserve some structure
                    text = soup.get_text(separator=' ', strip=True)
                    
                    # Also extract important metadata
                    title = soup.find('title')
                    if title:
                        text = f"Title: {title.string}\n\n{text}"
                    
                    return text
                    
            except ImportError:
                pass
            
            # Fallback to basic HTML parsing
            import html.parser
            
            class HTMLTextExtractor(html.parser.HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text_parts = []
                    self.skip_tags = {'script', 'style'}
                    self.current_tag = None
                
                def handle_starttag(self, tag, attrs):
                    self.current_tag = tag
                
                def handle_endtag(self, tag):
                    self.current_tag = None
                
                def handle_data(self, data):
                    if self.current_tag not in self.skip_tags:
                        text = data.strip()
                        if text:
                            self.text_parts.append(text)
            
            with open(file_path, 'r', encoding='utf-8') as f:
                parser = HTMLTextExtractor()
                parser.feed(f.read())
                return ' '.join(parser.text_parts)
                
        except Exception as e:
            self.logger.error(f"Failed to extract text from HTML {file_path}: {e}")
            return f"[Failed to extract HTML content from {file_path.name}: {str(e)}]"
    
    def _extract_structured_data(self, file_path: Path, file_type: FileType) -> Dict[str, Any]:
        """Extract structured data from file"""
        data = {}
        
        try:
            if file_type == FileType.JSON:
                import json
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
            # Add other structured data extraction as needed
        except Exception as e:
            self.logger.debug(f"Structured data extraction failed for {file_path}: {e}")
        
        return data
    
    def _get_image_dimensions(self, file_path: Path) -> Optional[tuple]:
        """Get image dimensions"""
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                return img.size  # Returns (width, height)
        except ImportError:
            self.logger.debug("Pillow not installed. Cannot get image dimensions.")
            return None
        except Exception as e:
            self.logger.debug(f"Failed to get image dimensions for {file_path}: {e}")
            return None
    
    def _get_document_page_count(self, file_path: Path) -> Optional[int]:
        """Get document page count"""
        try:
            if file_path.suffix.lower() == '.pdf':
                # Try PyPDF2
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        return len(reader.pages)
                except ImportError:
                    pass
                
                # Try PyMuPDF
                try:
                    import fitz
                    doc = fitz.open(file_path)
                    page_count = len(doc)
                    doc.close()
                    return page_count
                except ImportError:
                    pass
                    
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    # Approximate page count based on content
                    # This is not exact but gives an estimate
                    total_paragraphs = len(doc.paragraphs)
                    return max(1, total_paragraphs // 25)  # Rough estimate
                except ImportError:
                    pass
                    
        except Exception as e:
            self.logger.debug(f"Failed to get page count for {file_path}: {e}")
        
        return None
    
    def _create_artifacts(self, file_path: Path, metadata: FileMetadata) -> List[str]:
        """Create processing artifacts"""
        artifacts = []
        
        try:
            # Create metadata file
            metadata_file = self.output_dir / f"{file_path.stem}_metadata.json"
            import json
            with open(metadata_file, "w") as f:
                json.dump({
                    "file_path": metadata.file_path,
                    "file_name": metadata.file_name,
                    "file_size": metadata.file_size,
                    "file_type": metadata.file_type.value,
                    "hash_md5": metadata.hash_md5
                }, f, indent=2)
            artifacts.append(str(metadata_file))
            
            # Create thumbnail for images if enabled
            if (self.config.create_thumbnails and 
                metadata.file_type == FileType.IMAGE):
                thumbnail_path = self._create_thumbnail(file_path)
                if thumbnail_path:
                    artifacts.append(thumbnail_path)
                    
        except Exception as e:
            self.logger.debug(f"Artifact creation failed for {file_path}: {e}")
        
        return artifacts
    
    def _create_thumbnail(self, file_path: Path) -> Optional[str]:
        """Create image thumbnail"""
        try:
            from PIL import Image
            
            # Define thumbnail size
            thumbnail_size = (self.config.thumbnail_size, self.config.thumbnail_size)
            
            # Open and create thumbnail
            with Image.open(file_path) as img:
                # Convert RGBA to RGB if necessary
                if img.mode in ('RGBA', 'LA'):
                    # Create a white background
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                
                # Create thumbnail
                img.thumbnail(thumbnail_size, Image.Resampling.LANCZOS)
                
                # Save thumbnail
                thumbnail_path = self.output_dir / f"{file_path.stem}_thumbnail.jpg"
                img.save(thumbnail_path, 'JPEG', quality=85)
                
                return str(thumbnail_path)
                
        except ImportError:
            self.logger.debug("Pillow not installed. Cannot create thumbnails.")
            return None
        except Exception as e:
            self.logger.debug(f"Failed to create thumbnail for {file_path}: {e}")
            return None
    
    def get_processing_statistics(self, results: List[ProcessedFile]) -> Dict[str, Any]:
        """Calculate processing statistics"""
        total_files = len(results)
        completed = sum(1 for r in results if r.status == ProcessingStatus.COMPLETED)
        failed = sum(1 for r in results if r.status == ProcessingStatus.FAILED)
        
        # File type distribution
        type_counts = {}
        for result in results:
            file_type = result.metadata.file_type.value
            type_counts[file_type] = type_counts.get(file_type, 0) + 1
        
        # Size statistics
        total_size = sum(r.metadata.file_size for r in results)
        avg_size = total_size / total_files if total_files > 0 else 0
        
        # Processing time statistics
        total_time = sum(r.processing_time for r in results)
        avg_time = total_time / total_files if total_files > 0 else 0
        
        return {
            "total_files": total_files,
            "completed": completed,
            "failed": failed,
            "success_rate": completed / total_files if total_files > 0 else 0,
            "file_type_distribution": type_counts,
            "total_size_mb": total_size / (1024 * 1024),
            "average_file_size_kb": avg_size / 1024,
            "total_processing_time": total_time,
            "average_processing_time": avg_time
        }


# Factory function
def create_file_processor(**config_kwargs) -> FileProcessor:
    """Factory function to create file processor"""
    config = FileProcessorConfig(**config_kwargs)
    return FileProcessor(config)