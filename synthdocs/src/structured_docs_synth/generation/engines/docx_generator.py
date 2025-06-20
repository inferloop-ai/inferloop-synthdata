"""
DOCX generator for structured document generation
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

from ...core import (
    get_logger,
    get_config,
    DocumentGenerationError,
    ValidationError,
    get_document_type_config
)
from .template_engine import get_template_engine


class DOCXGenerator:
    """DOCX document generator using python-docx"""
    
    def __init__(self, output_dir: Optional[str] = None):
        self.logger = get_logger(__name__)
        self.config = get_config()
        
        # Set up output directory
        if output_dir:
            self.output_dir = Path(output_dir)
        else:
            self.output_dir = Path(self.config.generation.output_dir) / "docx"
        
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize template engine
        self.template_engine = get_template_engine()
        
        self.logger.info(f"DOCX generator initialized with output directory: {self.output_dir}")
    
    def generate_docx(
        self,
        document_type: str,
        data: Dict[str, Any],
        output_filename: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """Generate DOCX document from template and data"""
        
        try:
            # Render content from template
            content = self.template_engine.render_template(document_type, data)
            
            # Generate filename if not provided
            if not output_filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_filename = f"{document_type}_{timestamp}.docx"
            
            output_path = self.output_dir / output_filename
            
            # Create DOCX document
            doc = Document()
            
            # Set up document styles
            self._setup_document_styles(doc)
            
            # Add document properties
            self._set_document_properties(doc, document_type, metadata)
            
            # Build document content
            self._build_document_content(doc, document_type, content, data, metadata)
            
            # Save document
            doc.save(str(output_path))
            
            self.logger.info(f"Generated DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            raise DocumentGenerationError(
                f"Failed to generate DOCX for {document_type}: {str(e)}",
                details={
                    'document_type': document_type,
                    'output_filename': output_filename,
                    'error': str(e)
                }
            )
    
    def _setup_document_styles(self, doc: Document):
        """Set up custom styles for the document"""
        
        styles = doc.styles
        
        # Create custom heading style
        try:
            heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_style.paragraph_format.space_after = Pt(12)
            heading_style.paragraph_format.space_before = Pt(12)
        except ValueError:
            # Style already exists
            heading_style = styles['CustomHeading']
        
        # Create custom body style
        try:
            body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Arial'
            body_font.size = Pt(11)
            body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            body_style.paragraph_format.space_after = Pt(6)
        except ValueError:
            # Style already exists
            body_style = styles['CustomBody']
        
        # Create field label style
        try:
            label_style = styles.add_style('FieldLabel', WD_STYLE_TYPE.PARAGRAPH)
            label_font = label_style.font
            label_font.name = 'Arial'
            label_font.size = Pt(10)
            label_font.bold = True
            label_style.paragraph_format.space_after = Pt(3)
        except ValueError:
            # Style already exists
            label_style = styles['FieldLabel']
        
        # Create field value style
        try:
            value_style = styles.add_style('FieldValue', WD_STYLE_TYPE.PARAGRAPH)
            value_font = value_style.font
            value_font.name = 'Arial'
            value_font.size = Pt(11)
            value_style.paragraph_format.space_after = Pt(8)
            value_style.paragraph_format.left_indent = Inches(0.25)
        except ValueError:
            # Style already exists
            value_style = styles['FieldValue']
    
    def _set_document_properties(self, doc: Document, document_type: str, metadata: Optional[Dict[str, Any]]):
        """Set document properties and metadata"""
        
        core_props = doc.core_properties
        
        # Set basic properties
        core_props.author = "InferLoop SynthData Generator"
        core_props.title = metadata.get('title', f"{document_type.replace('_', ' ').title()} Document") if metadata else f"{document_type.replace('_', ' ').title()} Document"
        core_props.subject = f"Synthetic {document_type.replace('_', ' ')} document"
        core_props.created = datetime.now()
        core_props.modified = datetime.now()
        
        # Set additional properties from metadata
        if metadata:
            if 'description' in metadata:
                core_props.comments = metadata['description']
            if 'keywords' in metadata:
                core_props.keywords = metadata['keywords']
    
    def _build_document_content(
        self,
        doc: Document,
        document_type: str,
        content: str,
        data: Dict[str, Any],
        metadata: Optional[Dict[str, Any]] = None
    ):
        """Build the document content"""
        
        # Add document header if metadata provided
        if metadata and 'title' in metadata:
            title_para = doc.add_paragraph(metadata['title'])
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            doc.add_paragraph()  # Empty paragraph for spacing
        
        # Add document info table if metadata provided
        if metadata:
            info_items = []
            if 'date' in metadata:
                info_items.append(('Date:', metadata['date']))
            if 'reference' in metadata:
                info_items.append(('Reference:', metadata['reference']))
            if 'document_id' in metadata:
                info_items.append(('Document ID:', metadata['document_id']))
            
            if info_items:
                table = doc.add_table(rows=len(info_items), cols=2)
                table.style = 'Table Grid'
                
                for i, (label, value) in enumerate(info_items):
                    cells = table.rows[i].cells
                    cells[0].text = label
                    cells[1].text = str(value)
                    
                    # Format cells
                    cells[0].paragraphs[0].runs[0].font.bold = True
                    cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                    cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                
                doc.add_paragraph()  # Empty paragraph for spacing
        
        # Parse content into sections
        sections = self._parse_content_sections(content)
        
        # Build each section
        for section in sections:
            # Section title
            if section['title']:
                heading_para = doc.add_paragraph(section['title'])
                heading_para.style = 'CustomHeading'
            
            # Section content
            for line in section['content']:
                if line.strip():
                    # Check if it's a form field (contains ':')
                    if ':' in line and not line.strip().endswith(':'):
                        parts = line.split(':', 1)
                        label = parts[0].strip()
                        value = parts[1].strip()
                        
                        # Add field as separate paragraphs
                        label_para = doc.add_paragraph(f"{label}:")
                        label_para.style = 'FieldLabel'
                        
                        value_para = doc.add_paragraph(value)
                        value_para.style = 'FieldValue'
                    else:
                        # Regular paragraph
                        para = doc.add_paragraph(line)
                        para.style = 'CustomBody'
                else:
                    # Empty line - add empty paragraph for spacing
                    doc.add_paragraph()
        
        # Add signature section for certain document types
        if document_type in ['legal_contract', 'loan_application']:
            doc.add_page_break()
            self._add_signature_section(doc, document_type, data)
    
    def _parse_content_sections(self, content: str) -> List[Dict[str, Any]]:
        """Parse content into sections based on uppercase headers"""
        
        sections = []
        current_section = {'title': '', 'content': []}
        
        lines = content.split('\n')
        
        for line in lines:
            # Check if line is a section header (all uppercase)
            if line.strip() and line.strip().isupper() and len(line.strip()) > 3:
                # Save previous section if it has content
                if current_section['content']:
                    sections.append(current_section)
                # Start new section
                current_section = {'title': line.strip(), 'content': []}
            else:
                # Add line to current section
                current_section['content'].append(line)
        
        # Add last section
        if current_section['content'] or current_section['title']:
            sections.append(current_section)
        
        return sections
    
    def _add_signature_section(self, doc: Document, document_type: str, data: Dict[str, Any]):
        """Add signature section for documents that require it"""
        
        # Add signature heading
        signature_heading = doc.add_paragraph('SIGNATURES')
        signature_heading.style = 'CustomHeading'
        
        doc.add_paragraph()  # Empty paragraph for spacing
        
        # Determine parties for signatures
        if document_type == 'legal_contract' and 'parties' in data:
            parties = data['parties']
        else:
            parties = ['Party 1', 'Party 2']
        
        # Create signature table
        table = doc.add_table(rows=len(parties), cols=4)
        table.style = 'Table Grid'
        
        # Set column widths
        widths = [Inches(1.5), Inches(2.5), Inches(0.8), Inches(1.5)]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width
        
        # Fill table with signature lines
        for i, party in enumerate(parties[:2]):  # Max 2 parties for now
            cells = table.rows[i].cells
            
            # Party name
            cells[0].text = f"{party}:"
            cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Signature line
            cells[1].text = "_" * 40
            
            # Date label
            cells[2].text = "Date:"
            cells[2].paragraphs[0].runs[0].font.bold = True
            
            # Date line
            cells[3].text = "_" * 20
            
            # Format all cells
            for cell in cells:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def generate_batch(
        self,
        document_type: str,
        data_list: List[Dict[str, Any]],
        output_prefix: str = "",
        **kwargs
    ) -> List[Path]:
        """Generate multiple DOCX documents in batch"""
        
        generated_files = []
        
        for i, data in enumerate(data_list):
            try:
                filename = f"{output_prefix}{document_type}_{i+1:04d}.docx"
                output_path = self.generate_docx(
                    document_type,
                    data,
                    output_filename=filename,
                    **kwargs
                )
                generated_files.append(output_path)
                
            except Exception as e:
                self.logger.error(f"Failed to generate DOCX {i+1}/{len(data_list)}: {str(e)}")
                if self.config.generation.batch_continue_on_error:
                    continue
                else:
                    raise
        
        self.logger.info(f"Generated {len(generated_files)} DOCX documents")
        return generated_files